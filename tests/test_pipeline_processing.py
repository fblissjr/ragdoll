# tests/test_pipeline_processing.py
import pytest
import os
import json
from pathlib import Path
import numpy as np

# Assuming your src directory is in PYTHONPATH or tests are run from project root
# If src is a package and tests is a sibling, you might need to adjust imports
# e.g., from ..src import pipeline_orchestrator, ragdoll_config, ragdoll_utils, vector_store_manager
# For now, assuming direct import works due to PYTHONPATH setup or running pytest from root.
from src import pipeline_orchestrator
from src import ragdoll_config
from src import ragdoll_utils # For generate_display_source_name if checking metadata
from src import vector_store_manager # To load and check the Vicinity store
from vicinity import Metric # For type checking

# --- Test Helper Functions ---
def check_common_output_files(output_dir: Path, expect_viz: bool = False):
    """Checks for the presence of standard output files."""
    assert (output_dir / ragdoll_config.TEXT_CHUNKS_FILENAME).exists()
    assert (output_dir / ragdoll_config.CHUNK_IDS_FILENAME).exists()
    assert (output_dir / ragdoll_config.DETAILED_CHUNK_METADATA_FILENAME).exists()
    assert (output_dir / ragdoll_config.CHUNK_VECTORS_FILENAME).exists()
    
    vicinity_store_path = output_dir / ragdoll_config.VECTOR_STORE_SUBDIR_NAME
    assert vicinity_store_path.exists()
    assert (vicinity_store_path / "data.json").exists() # Vicinity internal file
    assert (vicinity_store_path / "vectors.npy").exists() # Vicinity internal file
    assert (vicinity_store_path / "arguments.json").exists() # Vicinity internal file

    if expect_viz:
        assert (output_dir / ragdoll_config.VISUALIZATION_DATA_FILENAME).exists()
    else:
        assert not (output_dir / ragdoll_config.VISUALIZATION_DATA_FILENAME).exists()

def load_and_validate_json(filepath: Path) -> Any:
    """Loads a JSON file and returns its content, failing test if error."""
    assert filepath.exists(), f"JSON file not found: {filepath}"
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data
    except json.JSONDecodeError as e:
        pytest.fail(f"Failed to decode JSON from {filepath}: {e}")
    except Exception as e:
        pytest.fail(f"Error loading JSON from {filepath}: {e}")


# --- Test Cases ---

def test_pipeline_default_chonkie_sdpm(temp_docs_dir, temp_output_dir):
    """
    Tests the pipeline with the default chunker (chonkie_sdpm), no classification, no viz.
    """
    pipeline_config = {
        "docs_folder": str(temp_docs_dir),
        "vector_data_dir": str(temp_output_dir),
        "overwrite": True,
        "gpu_device_id": -1, # Force CPU for tests for consistency
        "verbose": False,
        "embedding_model_name": ragdoll_config.DEFAULT_PIPELINE_EMBEDDING_MODEL, # Use a fast, small model if possible for tests
        "chunker_type": "chonkie_sdpm", # Explicitly default
        "num_processing_workers": 1, # Sequential for easier debugging if needed
        "enable_classification": False,
        "prepare_viz_data": False,
        # Add other necessary default params if not covered by _get_specific_chunker_params logic
        # For chonkie_sdpm, some defaults come from ragdoll_config.CHUNKER_DEFAULTS
        "chonkie_embedding_model": ragdoll_config.CHUNKER_DEFAULTS["chonkie_sdpm"]["embedding_model"],
        "chonkie_target_chunk_size": ragdoll_config.CHUNKER_DEFAULTS["chonkie_sdpm"]["chunk_size"],
        "chonkie_similarity_threshold": ragdoll_config.CHUNKER_DEFAULTS["chonkie_sdpm"]["threshold"],
        # ... ensure all required params for _get_specific_chunker_params are here or in defaults
    }
    
    success = pipeline_orchestrator.run_full_processing_pipeline(pipeline_config)
    assert success, "Pipeline run failed for chonkie_sdpm"
    
    check_common_output_files(temp_output_dir, expect_viz=False)

    # Validate content of some files
    text_chunks = load_and_validate_json(temp_output_dir / ragdoll_config.TEXT_CHUNKS_FILENAME)
    chunk_ids = load_and_validate_json(temp_output_dir / ragdoll_config.CHUNK_IDS_FILENAME)
    detailed_metadata = load_and_validate_json(temp_output_dir / ragdoll_config.DETAILED_CHUNK_METADATA_FILENAME)

    assert len(text_chunks) > 0, "No text chunks were generated."
    assert len(text_chunks) == len(chunk_ids) == len(detailed_metadata), "Mismatch in counts of chunks, IDs, and metadata."

    for meta_item in detailed_metadata:
        assert "vicinity_item_id" in meta_item
        assert "source_file" in meta_item
        assert "display_source_name" in meta_item
        # Check for some content from our sample files if possible (more advanced check)
        if "sample.txt" in meta_item["source_file"]:
             # Find the corresponding chunk text
             chunk_idx = chunk_ids.index(meta_item["vicinity_item_id"])
             assert "RAGdoll pipeline" in text_chunks[chunk_idx]


def test_pipeline_chonkie_recursive_markdown(temp_docs_dir, temp_output_dir):
    """
    Tests the pipeline with chonkie_recursive, assuming DOCX was converted to MD.
    Enable classification and visualization.
    """
    # Create a dummy sample.docx for this test to ensure Pandoc path is tested
    # This would ideally be a real, small DOCX file copied by conftest.py
    if not (temp_docs_dir / "sample.docx").exists():
        try:
            # Attempt to create a minimal docx if python-docx is available in test env
            from docx import Document
            doc = Document()
            doc.add_heading("DOCX Test Heading", level=1)
            doc.add_paragraph("This is a paragraph in a DOCX file.")
            doc.add_paragraph("Another list item for the DOCX.")
            doc.save(temp_docs_dir / "sample.docx")
            print("Created dummy sample.docx for test_pipeline_chonkie_recursive_markdown")
        except ImportError:
            pytest.skip("python-docx not available, skipping DOCX part of this test. Add a real sample.docx for full testing.")
        except Exception as e:
            print(f"Could not create dummy docx: {e}")
            pytest.skip("Could not create dummy sample.docx. Add a real one for full testing.")


    pipeline_config = {
        "docs_folder": str(temp_docs_dir),
        "vector_data_dir": str(temp_output_dir),
        "overwrite": True,
        "gpu_device_id": -1,
        "verbose": True, # Enable verbose for more test output
        "embedding_model_name": "sentence-transformers/all-MiniLM-L6-v2", # A known fast SBERT model
        "chunker_type": "chonkie_recursive",
        # For chonkie_recursive, specify a markdown ruleset if content is expected to be markdown
        "chonkie_basic_tokenizer": "ragdoll_utils.BGE_TOKENIZER_INSTANCE", # Use our tokenizer
        "chonkie_basic_chunk_size": 128, # Smaller chunk size for testing
         # Orchestrator's _get_specific_chunker_params will fetch "rules" and "lang" from CHUNKER_DEFAULTS for chonkie_recursive
         # We could override them here if needed, e.g. "chonkie_recursive_rules": "markdown"
        "num_processing_workers": 1,
        "enable_classification": True,
        "candidate_labels": ["technology", "general text", "list item"], # Simple labels for testing
        "classification_batch_size": 2,
        "prepare_viz_data": True,
        "viz_output_file": "test_viz_data.json",
        "umap_neighbors": 2, # Small value for small test dataset (ensure >1 if more than 2 chunks)
        "umap_min_dist": 0.01,
    }
    
    success = pipeline_orchestrator.run_full_processing_pipeline(pipeline_config)
    assert success, "Pipeline run failed for chonkie_recursive with classification/viz"

    check_common_output_files(temp_output_dir, expect_viz=True)
    assert (temp_output_dir / "test_viz_data.json").exists()

    detailed_metadata = load_and_validate_json(temp_output_dir / ragdoll_config.DETAILED_CHUNK_METADATA_FILENAME)
    text_chunks = load_and_validate_json(temp_output_dir / ragdoll_config.TEXT_CHUNKS_FILENAME)
    chunk_ids = load_and_validate_json(temp_output_dir / ragdoll_config.CHUNK_IDS_FILENAME)

    assert len(text_chunks) > 0
    found_docx_markdown_chunk = False
    found_md_chunk = False

    for i, meta_item in enumerate(detailed_metadata):
        assert "vicinity_item_id" in meta_item
        assert "source_file" in meta_item
        assert "classification_status" in meta_item or "top_label" in meta_item # Check for classification output
        if meta_item["classification_status"] != "skipped_classifier_init_failed": # if classifier ran
             assert "top_label" in meta_item
             assert "top_label_score" in meta_item
        if "sample.docx" in meta_item["source_file"]:
            assert meta_item.get("content_type") == "markdown", "DOCX should have been converted to markdown"
            # Check if chunking happened for the DOCX content
            if "DOCX Test Heading" in text_chunks[i] or "paragraph in a DOCX" in text_chunks[i]:
                found_docx_markdown_chunk = True
        if "sample.md" in meta_item["source_file"]:
            assert meta_item.get("content_type") == "markdown"
            if "Markdown Sample" in text_chunks[i] or "Item 1" in text_chunks[i]:
                found_md_chunk = True
    
    # These assertions depend on having actual sample.docx and sample.md processed
    # and that their content is distinctive enough to be found in chunks.
    if (temp_docs_dir / "sample.docx").exists():
        assert found_docx_markdown_chunk, "No chunks found that seem to originate from sample.docx content"
    if (temp_docs_dir / "sample.md").exists():
        assert found_md_chunk, "No chunks found that seem to originate from sample.md content"

    # Check visualization data
    viz_data = load_and_validate_json(temp_output_dir / "test_viz_data.json")
    assert isinstance(viz_data, list)
    if len(text_chunks) > 1: # UMAP runs if more than 1 point
      assert len(viz_data) == len(text_chunks)
      if viz_data:
          assert "id" in viz_data[0] and "x" in viz_data[0] and "y" in viz_data[0]
    elif len(text_chunks) == 1: # UMAP might not run or produce meaningful output for 1 point
        assert len(viz_data) == 0 or len(viz_data) == 1 # Depending on UMAP behavior with 1 point

    # Basic check of Vicinity store loading and metadata
    vsm = vector_store_manager.VectorStoreManager(str(temp_output_dir))
    assert vsm.load_store(), "Failed to load the created Vicinity store"
    assert vsm.get_total_items() == len(chunk_ids)
    store_meta = vsm.get_store_metadata()
    assert store_meta is not None
    assert store_meta.get("chunker_type") == "chonkie_recursive"
    assert store_meta.get("embedding_model_name") == "sentence-transformers/all-MiniLM-L6-v2"
    assert store_meta.get("classification_enabled") == True

# TODO: Add more tests:
# - Different Chonkie chunker types (chonkie_sentence, chonkie_token, chonkie_neural if feasible for quick tests)
# - Test with an empty docs_input_folder (should handle gracefully).
# - Test with --overwrite False when data already exists.
# - Test with different --num-processing-workers (e.g., 0 or >1, though 0 might default to 1 or max).
# - Test specific error conditions if possible (e.g., unreadable file, though this is hard to mock reliably for all file types).